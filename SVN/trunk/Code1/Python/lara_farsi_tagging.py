
# Tagging Farsi text using the hazm package

# So far, this has to be run under Python 2.7 because hazm includes wapiti, which doesn't tun under 3.7,
# so we can't use lara_utils.

from __future__ import unicode_literals
from hazm import *
 
lemmatizer = Lemmatizer()

def tag_file(infile, outfile):
    instring = read_lara_text_file(infile)
    if not instring:
        return False
    outstring = tag_string(instring)
    if not outstring:
        print('*** Error: unable to tag string')
        return False
    write_lara_text_file(outstring, outfile)
    print('--- Call to Farsi tagger successfully completed')
    return True

def tag_string(string):
    tokens = word_tokenize(string)
    lemmatized_tokens = [ token_to_tag_pair(token) for token in tokens ]
    return substitute_serial(lemmatized_tokens, string, '')

# If we've got a verb, the lemma is of the form PastStem#PresentStem.
# The tag should be formed by adding an 'n' ('\u0646') to PastStem.
def token_to_tag_pair(token):
    normalised_token = normalise_token(token)
    lemma0 = lemmatizer.lemmatize(token)
    lemma = lemma0.split('#')[0] + '\u0646' if '#' in lemma0 else lemma0
    if normalised_token == lemma:
        tagged_token = normalised_token
    else:
        tagged_token = normalised_token + '#' + lemma + '#'
        #print('-- Tagged ' + token + ' as ' + tagged_token)
    return (normalised_token, tagged_token)

# Sometimes a word gets connected to the following word with an underscore.
#
# The lemmatizer appears to want this form, but we need to remove the underscore
# to retrieve the original orthography
def normalise_token(item):
    return item.replace('_', ' ')

def substitute_serial(substList, string, separator):
    instring = string
    outstring = ''
    for subst in substList:
        ( w, w1 ) = subst
        index = instring.find(w)
        if index == -1:
            print('--- Unable to find "{w}" in "{string}"'.format(w=w, string=string))
            return False
        else:
            outstring = outstring + instring[:index] + w1 + separator
            instring = instring[index + len(w):]
    return outstring + instring

# ----------------------------------
# Utils

# These are messed-around versions of functions in lara_utils.py, adapted for 2.7

import os.path
# Use io package to get a version of 'open' with the encoding parameter that works in 2.7
import io

def absolute_file_name(pathname):
    return os.path.abspath(os.path.expandvars(pathname)).replace('\\', '/')

def file_exists(pathname):
    return os.path.isfile(absolute_file_name(pathname))

## We put iso-8859-1 earlier in lara_utils.py, but here it's very unlikely to be needed
encodings_to_try = ['utf-8-sig', 'utf-8', 'utf-16le', 'utf-16be', 'iso-8859-1']

## Read a LARA unicode text file and return a string.
## Try some other encodings if utf-8 doesn't work.
def read_lara_text_file(pathname):
    global encodings_to_try
    if not file_exists(pathname):
        print('*** Error: unable to find ' + pathname)
        return False
    elif extension_for_file(pathname) == 'docx':
        return read_docx_file(pathname)
    else:
        abspathname = absolute_file_name(pathname)
        return read_lara_text_file1(abspathname, encodings_to_try)

def extension_for_file(pathname):
    Components = absolute_file_name(pathname).split('.')
    N = len(Components)
    if N == 1:
        return ''
    else:
        return Components[N-1] 

def read_lara_text_file1(abspathname, encodings):
    if len(encodings) == 0:
        print('*** Error: no more encodings to try, giving up.')
        return False
    else:
        ( this_encoding, remaining_encodings ) = ( encodings[0], encodings[1:] )
        try:
            with io.open(abspathname, encoding=this_encoding) as f:
                lines = [ line for line in f ]
            print('--- Read LARA text file as ' + this_encoding + ' ' + str(len(lines)) + ' lines ' + abspathname)
            return "".join(lines)
        except Exception as e:
            print('*** Warning: unable to read text file ' + abspathname + ' as ' + this_encoding)
            print(str(e))
            return read_lara_text_file1(abspathname, remaining_encodings)

## Read a .docx file and turn it into a string

def read_docx_file(pathname):
    try:
        import docx
        abspathname = absolute_file_name(pathname)
        Doc = docx.Document(abspathname)
        Lines = [ Paragraph.text for Paragraph in Doc.paragraphs ]
        return '\n'.join(Lines)
    except Exception as e:
            print('*** Error when trying to read .docx file ' + pathname)
            print(str(e))
            return False

## Write a string to a unicode text file, printing message
def write_lara_text_file(string, pathname):
    try:
        abspathname = absolute_file_name(pathname)
        if extension_for_file(abspathname) == 'docx':
            write_docx_file(string, abspathname)
        else:
            with io.open(abspathname, 'w', encoding='utf-8') as f:
                f.write(string)
            print('--- Written LARA text file ' + abspathname)
    except Exception as e:
        print('*** Error: when trying to write LARA text file ' + pathname)
        print(str(e))
        return False

## Write a string to a .docx file

def write_docx_file(string, pathname):
    try:
        import docx
        abspathname = absolute_file_name(pathname)
        Doc = docx.Document()
        Lines = string.split('\n')
        for Line in Lines:
            Doc.add_paragraph(Line)
        Doc.save(abspathname)
        print('--- Written LARA text file ' + abspathname)
        return True
    except Exception as e:
            print_and_flush('*** Error when trying to write .docx file ' + pathname)
            print_and_flush(str(e))
            return False

    
