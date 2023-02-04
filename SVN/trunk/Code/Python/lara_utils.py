
import os.path
import os
import stat
import time
import shutil
from pathlib import Path
import json
import csv
import requests
import zipfile
import sys
import gzip
import pickle
import pprint
from datetime import datetime
import re
import random
import SVN.trunk.Code.Python.wagnerfischer as wagnerfischer
#import  wagnerfischer
## Simple utility functions etc

# ----------------------------------

## Working directories

temp_files_and_directories = []

## Create a tmp file
def get_tmp_trace_file(Params):
    return get_tmp_file_with_extension(Params, 'txt')

def get_tmp_csv_file(Params):
    return get_tmp_file_with_extension(Params, 'csv')

def get_tmp_json_file(Params):
    return get_tmp_file_with_extension(Params, 'json')

def get_tmp_file_with_extension(Params, Extension):
    working_tmp_directory = Params.working_tmp_directory
    random.seed()
    I = random.randint(1, 10000000000000000)
    create_directory_if_it_doesnt_exist(working_tmp_directory)
    pathname = f'{working_tmp_directory}/lara_tmp_{I}.{Extension}'
    abspathname = absolute_file_name(pathname)
    if not file_exists(abspathname):
        temp_files_and_directories.append( abspathname )
        return abspathname
    else:
        return get_tmp_trace_file()

## Create a tmp directory 
def get_tmp_directory(Params):
    working_tmp_directory = Params.working_tmp_directory
    random.seed()
    I = random.randint(1, 10000000000000000)
    create_directory_if_it_doesnt_exist(working_tmp_directory)
    pathname = f'{working_tmp_directory}/lara_tmp_dir_{I}'
    abspathname = absolute_file_name(pathname)
    if not directory_exists(abspathname):
        create_directory(abspathname)
        temp_files_and_directories.append( abspathname )
        return abspathname
    else:
        return get_tmp_directory()

def cleanup_temp_files_and_directories():
    for abspathname in temp_files_and_directories:
        if os.path.isfile( abspathname ):
            delete_file_if_it_exists( abspathname )
        if os.path.isdir( abspathname ):
            delete_directory_if_it_exists( abspathname )

# ----------------------------------

_encodingError = '(... something in an encoding that can\'t be written out ...)'

def report_encoding_error(Str):
    try:
        import unidecode
        Str1 = unidecode.unidecode(str(Str))
        print(f'{Str1} [some characters changed, encoding problems]')
    except:
        print(_encodingError)

## Print, flushing output
def print_and_flush(Object):
    try:
        print(Object, flush=True)
    except:
        #print(_encodingError)
        report_encoding_error(Object)

## Print, flushing output, no newline
def print_and_flush_no_newline(Object):
    try:
        print(Object, end='', flush=True)
    except:
        #print(_encodingError)
        report_encoding_error(Object)

stored_timing_messages = []

def init_stored_timings():
    global stored_timing_messages
    stored_timing_messages = []

def print_stored_timings():
    global stored_timing_messages
    if len(stored_timing_messages) > 0:
        print_and_flush('\n--- All timings:')
        print_and_flush('\n'.join(stored_timing_messages))

## Print with elapsed time, flushing output
def print_and_flush_with_elapsed_time(Message, StartTime):
    global stored_timing_messages
    try:
        ElapsedTime = time.time() - StartTime
        MessageWithTiming = f'{Message} ({ElapsedTime:.3} secs)'
        print(MessageWithTiming, flush=True)
        stored_timing_messages += [ MessageWithTiming ]
    except:
        #print(_encodingError)
        report_encoding_error(MessageWithTiming)

all_warnings = {}

def init_warnings():
    global all_warnings
    all_warnings = {}

## Print warning if it hasn't already appeared, flushing output 
def print_and_flush_warning(Object):
    global all_warnings
    if not isinstance(Object, str) or Object in all_warnings:
        return
    all_warnings[Object] = 'presented'
    try:
        print(Object, flush=True)
    except:
        #print(_encodingError)
        report_encoding_error(Object)

## Check if file exists
def file_exists(pathname):
    return os.path.isfile(absolute_file_name(pathname))

def check_files_exist(Files):
    AllExist = True
    for File in Files:
        if not file_exists(File):
            print_and_flush(f'*** Error: file not found: {File}')
            AllExist = False
    return AllExist

## Check if directory exists
def directory_exists(pathname):
    return os.path.isdir(absolute_file_name(pathname))

## File1 and File2 both exist, ages of File1 less than age of File2
def file_is_newer_than_file(File1, File2):
    return file_exists(File1) and \
           file_exists(File2) and \
           file_age_in_seconds(File1) < file_age_in_seconds(File2)

## File1 exists, all files in Files exist and are older than File1
def file_is_newer_than_files(File1, Files):
    if not file_exists(File1):
        return False
    Age1 = file_age_in_seconds(File1)
    for File2 in Files:
        if not Age1 < file_age_in_seconds(File2):
            return False
    return True

## Get age of file
def file_age_in_seconds(pathname):
    return time.time() - os.stat(absolute_file_name(pathname))[stat.ST_MTIME]

## Get size of file
def file_size_in_bytes(pathname):
    return os.stat(absolute_file_name(pathname))[stat.ST_SIZE]

## Strip off extension and return base file and extension
def file_to_base_file_and_extension(pathname):
    Components = absolute_file_name(pathname).split('.')
    N = len(Components)
    if N == 1:
        return ( pathname, '' )
    else:
        return ( '.'.join(Components[:N-1]), Components[N-1].lower() )

## Change the extension in a pathname
def change_extension(pathname, extension):
    Components = pathname.split('.')
    N = len(Components)
    Base = pathname if N == 1 else '.'.join(Components[:N-1])
    return f'{Base}.{extension}'

## Return just the extension
def extension_for_file(pathname):
    return file_to_base_file_and_extension(pathname)[1]

## Return just the base name
def base_name_for_pathname(pathname):
    Components = absolute_file_name(pathname).split('/')
    return Components[-1]

## Return just the directory
def directory_for_pathname(pathname):
    Components = absolute_file_name(pathname).split('/')
    return '/'.join(Components[:-1])

## Age of file
def file_age_in_seconds(pathname):
    return time.time() - os.stat(absolute_file_name(pathname))[stat.ST_MTIME]

## Check that an environment variable exists and points to a directory
def check_environment_variable_directory(Var):
    if Var in os.environ and os.path.isdir(os.path.abspath(os.environ[Var])):
        return True
    elif not Var in os.environ:
        print_and_flush(f'*** Error: environment variable {Var} not found')
        return False
    else:
        print_and_flush(f'*** Error: environment variable {Var} not a directory')
        return False

## If a bad environment variable directory is referenced, print an error and return False
def check_environment_variable_directories_in_pathname(pathname):
    FirstComponent = pathname.replace('\\', '/').split('/')[0]
    if FirstComponent.startswith('$'):
        return check_environment_variable_directory(FirstComponent[1:])
    else:
        return True

## Expand environment vars etc to get an absolute path name
def absolute_file_name(pathname):
    if not isinstance(pathname, str):
        print_and_flush(f'*** Error: bad argument "{pathname}" to lara_utils.absolute_file_name ')
        return False
    if check_environment_variable_directories_in_pathname(pathname):
        return os.path.abspath(os.path.expandvars(pathname)).replace('\\', '/').replace('//', '/')
    else:
        return pathname + ' (bad environment variable)'

## Get absolute filename of a file in Python/Code directory
def absolute_filename_in_python_directory(basename):
    if hasattr(sys, "_MEIPASS"): # Pyinstaller?
        return os.path.join(sys._MEIPASS, basename).replace('\\', '/')
    import inspect
    codefilename = inspect.getframeinfo(inspect.currentframe()).filename
    codepath = os.path.dirname(os.path.abspath(codefilename))
    return os.path.join(codepath, basename).replace('\\', '/')

def rename_file( oldname, newname ):
    os.rename( absolute_file_name(oldname), absolute_file_name(newname) )
    # os.rename doesn't return any value :( 
    # Therefore we can only check for the existence and non-existence of the files
    if not file_exists(newname) or file_exists( oldname ):
        print_and_flush(f'*** Error: cannot rename file "{oldname} to ${newname}.')
        return False
    return True

def is_valid_contentid( id ):
    return id and id != ""
    ## import re
    ## return re.search( "^[a-z][a-z0-9_]+$", id, re.IGNORECASE )

## return the name of the current python executable ("python" or "python3")
def python_executable():
    import sys
    return os.path.splitext(os.path.basename(sys.executable))[0]

## Execute a command-line call and return the code
def execute_lara_os_call_direct(Command):
    Code = False
    try:
        Code = os.system(Command)
        if Code != 0:
            print_and_flush(f'*** Warning: direct OS call failed, code = {Code}: "{shorten_for_printing(Command)}"')
    except Exception as e:
        print_and_flush(f'*** Warning: something went wrong during direct OS call: "{shorten_for_printing(Command)}"')
        print_and_flush(str(e))
    return Code

## Make call through bash
def execute_lara_os_call(Command, Params):
    return safe_execute_with_bash(Command, Params)

def safe_execute_with_bash(Command, Params):
    if not bash_found():
        print_and_flush(f'*** Error: unable to find "bash" in $PATH, cannot make system call "Command"')
        return 1
    TmpFile = get_tmp_trace_file(Params)
    write_unicode_text_file(Command, TmpFile)
    Code = False
    try:
        Code = os.system(f'bash < {TmpFile}')
        if Code != 0:
            print_and_flush(f'*** Warning: bash call failed, code = {Code}: "{shorten_for_printing(Command)}"')
    except Exception as e:
        print_and_flush(f'*** Warning: something went wrong during bash call: "{shorten_for_printing(Command)}"')
        print_and_flush(str(e))
    finally:
        delete_file_if_it_exists(TmpFile)
    return Code

def bash_found():
    return shutil.which('bash') != None

## Make sure string is less than 2000 chars long
def shorten_for_printing(Str):
    return Str if len(Str) < 2000 else f'{Str[:980]} ( ... {len(Str) - 1960} chars skipped ...) {Str[-980:]})'

## Delete a file if it's there
def delete_file_if_it_exists(pathname):
    try:
        abspathname = absolute_file_name(pathname)
        if file_exists(abspathname):
            os.remove(abspathname)
            if file_exists(abspathname):
               print_and_flush(f'*** Error: unable to delete file {pathname}')
               return False
            else:
                return True
        else:
            return True
    except:
        print_and_flush(f'*** Error: unable to delete file {pathname}')
        return False

## Delete a directory if it's there
def delete_directory_if_it_exists(pathname):
    abspathname = absolute_file_name(pathname)
    try:
        if directory_exists(abspathname):
            Result = shutil.rmtree(abspathname)
            print_and_flush(f'--- Deleted directory {pathname}')
            return Result
        return True
    except:
        print_and_flush(f'*** Error: unable to delete directory {pathname}')
        return False

## Create a directory, making directories above if necessary
def create_directory(pathname):
    abspathname = absolute_file_name(pathname)
    try:
        if not directory_exists(f'{abspathname}/..'):
            create_directory(f'{abspathname}/..')
        os.mkdir(abspathname)
        print_and_flush(f'--- Created directory {pathname}')
        return True
    except:
        print_and_flush(f'*** Error: unable to create directory {pathname}')
        return False

## Create a directory, deleting it first if it already exists
def create_directory_deleting_old_if_necessary(pathname):
    if not delete_directory_if_it_exists_try_n_times(3, pathname):
        return False
    return create_directory_try_n_times(3, pathname)

def delete_directory_if_it_exists_try_n_times(n, pathname):
    if n < 1:
        print_and_flush(f'*** Error: giving up on trying to delete directory {pathname}')
        return False
    if delete_directory_if_it_exists(pathname):
        return True
    else:
        time.sleep(1)
        return delete_directory_if_it_exists_try_n_times(n-1, pathname)
        
def create_directory_try_n_times(n, pathname):
    if n < 1:
        print_and_flush(f'*** Error: giving up on trying to create directory {pathname}')
        return False
    if create_directory(pathname):
        return True
    else:
        time.sleep(1)
        return create_directory_try_n_times(n-1, pathname)

## Create a directory if it doesn't already exist
def create_directory_if_it_doesnt_exist(pathname):
    if not directory_exists(pathname):
        return create_directory(pathname)
    else:
        return True

## Get files in directory
def directory_files(pathname):
    abspathname = absolute_file_name(pathname)
    # It's convenient to say there are no files in a nonexistent directory
    return os.listdir(abspathname) if directory_exists(abspathname) else []

## Get files in directory, returning only the ones with the named extension
def files_with_given_extension_in_directory(pathname, extension):
    return files_with_one_of_given_extensions_in_directory(pathname, [ extension ])

def files_with_one_of_given_extensions_in_directory(pathname, extensions):
    return [ File for File in directory_files(pathname)
             if file_to_base_file_and_extension(File)[1] in extensions ]

## Subdirectories of directory
def directory_members_of_directory(Dir):
    return [ Subdir for Subdir in directory_files(Dir) if directory_exists(f'{Dir}/{Subdir}') ]

## File (i.e. nondirectory) members of directory
def file_members_of_directory(Dir):
    return [ File for File in directory_files(Dir) if file_exists(f'{Dir}/{File}') ]

## Move file or directory
def move(pathname1, pathname2):
    abspathname1 = absolute_file_name(pathname1)
    abspathname2 = absolute_file_name(pathname2)
    try:
        shutil.move(abspathname1, abspathname2)
        return True
    except:
        return False

## Copy file to file or directory
def copy_file(pathname1, pathname2):
    abspathname1 = absolute_file_name(pathname1)
    abspathname2 = absolute_file_name(pathname2)
    try:
        shutil.copy2(abspathname1, abspathname2)
        return True
    except Exception as e:
        print_and_flush(f'*** Error: failed to copy {pathname1} to {pathname2}')
        print_and_flush(str(e))
        return False

def copy_file_if_it_exists_else_warn(pathname1, pathname2):
    if not file_exists(pathname1):
        print_and_flush(f'*** Warning: file not found: {pathname1}')
        return True
    else:
        return copy_file(pathname1, pathname2)

def copy_file_with_feedback(pathname1, pathname2):
    abspathname1 = absolute_file_name(pathname1)
    abspathname2 = absolute_file_name(pathname2)
    try:
        shutil.copy2(abspathname1, abspathname2)
        if not file_exists(abspathname2):
            print_and_flush(f'*** Error: failed to copy {pathname1} to {pathname2}')
            return False
        return True
    except Exception as e:
        print_and_flush(f'*** Error: failed to copy {pathname1} to {pathname2}')
        print_and_flush(str(e))
        return False

## Return the number of file in Dir that have extension in the list Extensions
def number_of_directory_files_with_extension(Dir, Extensions):
    if not directory_exists(Dir):
        return 0
    return len([ File for File in directory_files(Dir)
                 if Extensions == 'all' or extension_for_file(File) in Extensions ])

## Copy files with given extensions from Dir to Dir1. If Extensions is 'all', copy all files.
def copy_directory_one_file_at_a_time(Dir, Dir1, Extensions):
    if not directory_exists(Dir):
        print_and_flush(f'*** Warning: source dir {Dir} not found, copy_directory_one_file_at_a_time has nothing to do')
        return
    if not directory_exists(Dir1):
        print_and_flush(f'*** Error: target dir {Dir1} not found, copy_directory_one_file_at_a_time failed')
        return
    Failed = 0
    Copied = 0
    Files = directory_files(Dir)
    if len(Files) > 0:
        for File in Files:
            if Extensions == 'all' or extension_for_file(File) in Extensions:
                Result = copy_file(f'{Dir}/{File}', f'{Dir1}/{File}')
                if not Result:
                    print_and_flush(f'*** Error: unable to copy file {Dir}/{File} to {Dir1}')
                    Failed += 1
                else:
                    Copied += 1
        if Failed > 0:
            print_and_flush(f'*** Error: {Failed} files out of {Failed + Copied} not copied')
        else:
            print_and_flush(f'--- Copied {Copied} files')

def copy_named_files_between_directories(Files, Dir, Dir1):
    Failed = 0
    Copied = 0
    for File in Files:
        ( FromFile, ToFile ) = ( f'{Dir}/{File}', f'{Dir1}/{File}' )
        if not file_exists(FromFile):
            print_and_flush(f'*** Error: file not found {FromFile}')
            Failed += 1
        else:
            Result = copy_file(FromFile, ToFile)
            if not Result:
                print_and_flush(f'*** Error: unable to copy file {Dir}/{File} to {Dir1}')
                Failed += 1
            else:
                Copied += 1
    if Failed > 0:
        print_and_flush(f'*** Error: {Failed} files out of {Failed + Copied} not copied')
    else:
        print_and_flush(f'--- Copied {Copied} files')
    
## Make directory or file writable

def make_writable(pathname):
     os.chmod(absolute_file_name(pathname), 777)

encodings_to_try = ['utf-8-sig', 'utf-8', 'iso-8859-1', 'utf-16le', 'utf-16be']

## Read a LARA text file and return a string.
## If it's a .docx extension, read it as a .docx file
## Then try to read it as a utf-8 encoded file
## Try some other encodings if utf-8 doesn't work.
def read_lara_text_file(pathname):
    global encodings_to_try
    if not file_exists(pathname):
        print_and_flush(f'*** Error: unable to find {pathname}.')
        return False
    elif extension_for_file(pathname) == 'docx':
        return read_docx_file(pathname)
    else:
        abspathname = absolute_file_name(pathname)
        return read_lara_text_file1(abspathname, encodings_to_try)

def read_lara_text_file1(abspathname, encodings):
    if len(encodings) == 0:
        print_and_flush(f'*** Error: no more encodings to try, giving up.')
        return False
    else:
        ( this_encoding, remaining_encodings ) = ( encodings[0], encodings[1:] )
        try:
            with open(abspathname, encoding=this_encoding) as f:
                lines = [ line for line in f ]
            print_and_flush(f'--- Read LARA text file as {this_encoding} ({len(lines)} lines) {abspathname}')
            return "".join(lines)
        except Exception as e:
            print_and_flush(f'*** Warning: unable to read text file {abspathname} as {this_encoding}')
            print_and_flush(str(e))
            return read_lara_text_file1(abspathname, remaining_encodings)

## Read a .docx file and turn it into a string

def read_docx_file(pathname):
    #Text1 = read_docx_file_with_docx2txt(pathname)
    #if Text1:
    #    return Text1
    Text2 = read_docx_file_with_docx(pathname)
    if Text2:
        return Text2
    print_and_flush(f'*** Error when trying to read .docx file {pathname}')

## The advantage of the docx2txt package (https://github.com/ankushshah89/python-docx2txt)
## is that it handles links.
##
## (Update, 20191031: but unfortunately it also introduces random tab chars)
def read_docx_file_with_docx2txt(pathname):
    try:
        import docx2txt
        abspathname = absolute_file_name(pathname)
        return docx2txt.process(abspathname).replace('\n\n', '\n')
    except Exception as e:
            print_and_flush(f'*** Warning: error when trying to read .docx file with docx2txt')
            print_and_flush(str(e))
            return False

## If that's not installed, maybe docx is?
def read_docx_file_with_docx(pathname):
    try:
        import docx
        abspathname = absolute_file_name(pathname)
        Doc = docx.Document(abspathname)
        Lines = [ Paragraph.text for Paragraph in Doc.paragraphs ]
        return '\n'.join(Lines)
    except Exception as e:
            print_and_flush(f'*** Warning: error when trying to read .docx file with docx')
            print_and_flush(str(e))
            return False

## Read UTF-8 encoded tab separated CSV to a list of lists
## Try some other encodings if utf-8 doesn't work.
def read_lara_csv(pathname):
    return read_lara_csv_specifying_quoting(pathname, 'noquotes')
        
def read_lara_csv_specifying_quoting(pathname, Quotes):
    global encodings_to_try
    abspathname = absolute_file_name(pathname)
    if not file_exists(abspathname):
        print_and_flush(f'*** Error: unable to find {pathname}')
        return False
    else:
        return read_lara_csv1(abspathname, encodings_to_try, Quotes)
    # There have been some weird problems with file_exists apparently not always returning
    # the right result. Alternately try skipping the test here and relying on the try/except
    # to catch missing files instead, but it doesn't seem to make any difference.
    #read_lara_csv1(abspathname, encodings_to_try, Quotes)

def read_lara_csv1(abspathname, encodings, Quotes):
    if len(encodings) == 0:
        print_and_flush(f'*** Error: no more encodings to try, giving up.')
        return False
    else:
        ( this_encoding, remaining_encodings ) = ( encodings[0], encodings[1:] )
        try:
            with open(abspathname, 'r', encoding=this_encoding) as f:
                if Quotes == 'quotes':
                    reader = csv.reader(f, delimiter='\t', quotechar='"')
                else:
                    reader = csv.reader(f, delimiter='\t', quotechar='"', quoting=csv.QUOTE_MINIMAL)
                List = list(reader)
            f.close()
            print_and_flush(f'--- Read CSV spreadsheet as {this_encoding} ({len(List)} lines) {abspathname}')
            return List
        except Exception as e:
            print_and_flush(f'*** Error: when trying to read CSV {abspathname} as {this_encoding}')
            print_and_flush(str(e))
            return read_lara_csv1(abspathname, remaining_encodings, Quotes)        

## Read an unicode text file and return a list of lines
def read_unicode_text_file_to_lines(pathname):
    abspathname = absolute_file_name(pathname)
    with open(abspathname, encoding='utf-8') as f:
        lines = [ line for line in f ]
    return lines

## Read an ascii text file and return a list of lines
def read_ascii_text_file_to_lines(pathname):
    abspathname = absolute_file_name(pathname)
    with open(abspathname) as f:
        lines = [ line for line in f ]
    return lines

## Write a string to a file, printing message
## If extension is .docx, try to write it out that way
## Else write it out as utf-8
def write_lara_text_file(string, pathname):
    try:
        abspathname = absolute_file_name(pathname)
        if extension_for_file(abspathname) == 'docx':
            write_docx_file(string, abspathname)
        else:
            with open(abspathname, 'w', encoding='utf-8') as f:
                f.write(string)
        print_and_flush(f'--- Written LARA text file {abspathname}')
        return True
    except Exception as e:
        print_and_flush(f'*** Error: when trying to write LARA text file {pathname}')
        print_and_flush(str(e))
        return False

## Write a string to a unicode text file, not printing message
def write_unicode_text_file(string, pathname):
    try:
        abspathname = absolute_file_name(pathname)
        with open(abspathname, 'w', encoding='utf-8') as f:
            f.write(string)
    except Exception as e:
        print_and_flush(f'*** Error: when trying to write LARA unicode text file {pathname}')
        print_and_flush(str(e))
        return False

## Write a string to a .docx file

def write_docx_file(string, pathname):
    try:
        import docx
        abspathname = absolute_file_name(pathname)
        Doc = docx.Document()
        Lines = string.split('\n')
        for Line in Lines:
            Doc.add_paragraph(Line)
        Doc.save(abspathname)
        return True
    except Exception as e:
            print_and_flush(f'*** Error when trying to write .docx file {pathname}')
            print_and_flush(str(e))
            return False

## Read a JSON file to a Python structure
def read_json_file(pathname):
    abspathname = absolute_file_name(pathname)
    try:
        with open(abspathname, encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print_and_flush(f'*** Error: unable to read JSON file {pathname}')
        print_and_flush(str(e))
        return False

## Write a Python structure to a JSON file
def write_json_to_file(data, pathname):
    abspathname = absolute_file_name(pathname)
    try:
        with open(abspathname, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=4, sort_keys=True) # prettyprint json output, and sort any keys
            # json.dump(data, f)
        print_and_flush(f'--- Written JSON file {pathname}')
        return True
    except Exception as e:
        print_and_flush(f'*** Error: unable to write JSON to file {pathname}')
        print_and_flush(str(e))
        return False

def write_json_to_file_plain_utf8(data, pathname):
    abspathname = absolute_file_name(pathname)
    try:
        with open(abspathname, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=4, sort_keys=True, ensure_ascii=False) # don't escape using \u
        print_and_flush(f'--- Written JSON file {pathname}')
        return True
    except Exception as e:
        print_and_flush(f'*** Error: unable to write JSON to file {pathname}')
        print_and_flush(str(e))
        return False

def prettyprint_json_file(infile, outfile):
    Data = read_json_file(infile)
    if not Data == False:
        write_json_to_file_plain_utf8(Data, outfile)

## Read a file from a URL, put it in a local file and return the name of that
## file or False if it didn't work
def read_file_from_url(url, pathname):
    try:
        abspathname = absolute_file_name(pathname)
        r = requests.get(url)
        if r.status_code == 200:
            if r.encoding:
                encoding = r.encoding
                text = (r.text).replace('\r', '')
                with open(abspathname, 'w', encoding=encoding) as f:
                    f.write(text)
                print_and_flush(f'--- Downloaded text file {abspathname} from {url}')
                return pathname
            else:
                with open(abspathname, 'wb') as f:
                    f.write(r.content)
                print_and_flush(f'--- Downloaded non-text file {abspathname} from {url}')
                return pathname
        else:
            return False
    except Exception as e:
        print_and_flush(f'*** Error: unable to download from {url}')
        print_and_flush(str(e))
        return False

## As above, but use streaming download
def read_file_from_url_streamed(url, pathname):
    try:
        abspathname = absolute_file_name(pathname)
        with requests.get(url, stream=True) as r:
            with open(abspathname, 'wb') as f:
                shutil.copyfileobj(r.raw, f)
        if r.status_code == 200 and file_exists(abspathname):
            print_and_flush(f'--- Downloaded file {pathname}')
            return pathname
        else:
            print_and_flush(f'*** Error: unable to download from {url}')
            return False
    except Exception as e:
        print_and_flush(f'*** Error: unable to download from {url}')
        print_and_flush(str(e))
        return False

## Read a file from a URL
def read_lines_from_url(url):
    r = requests.get(url)
    if r.status_code == 200:
        return r.text.split()
    else:
        return False

## Read JSON from a URL
def read_json_from_url(url):
    r = requests.get(url)
    if r.status_code == 200:
        try:
            JSON = r.json()
        except:
            print_and_flush(f'*** Warning: bad JSON at {url}')
            JSON = False
        return JSON
    else:
        return False

## Write list of lists to UTF-8 encoded tab separated CSV  
def write_lara_csv(Lines, pathname):
    abspathname = absolute_file_name(pathname)
    with open(abspathname, 'w', encoding="utf-8") as f:
        #writer = csv.writer(f, delimiter='\t', quotechar='"', lineterminator="\n")
        #writer = csv.writer(f, delimiter='\t', quoting=csv.QUOTE_NONE, escapechar='\\', lineterminator="\n")
        writer = csv.writer(f, delimiter='\t', quoting=csv.QUOTE_MINIMAL, escapechar='\\', quotechar='"', lineterminator="\n")
        writer.writerows(Lines)
        print_and_flush(f'--- Written CSV spreadsheet ({len(Lines)} lines) {pathname}')
    f.close()

def write_lara_comma_csv(Lines, pathname):
    abspathname = absolute_file_name(pathname)
    with open(abspathname, 'w', encoding="utf-8") as f:
        writer = csv.writer(f, delimiter=',', quoting=csv.QUOTE_MINIMAL, escapechar='\\', quotechar='"', lineterminator="\n")
        writer.writerows(Lines)
        print_and_flush(f'--- Written CSV spreadsheet with comma separators ({len(Lines)} lines) {pathname}')
    f.close()

## Extract a zipfile to a directory
def unzip_file(pathname, dir):
    abspathname = absolute_file_name(pathname)
    absdirname = absolute_file_name(dir)
    if not file_exists(abspathname):
        print_and_flush(f'*** Error: unable to find {pathname}')
        return False
    try:
        zip_ref = zipfile.ZipFile(abspathname, 'r')  
        zip_ref.extractall(absdirname)
        zip_ref.close()
        print_and_flush(f'--- Unzipped {pathname} to {dir}')
        return True
    except Exception as e:
        #print_and_flush(f'*** Error: something went wrong when trying to unzip {pathname} to {dir}')
        print_and_flush(f'*** Warning: something went wrong when trying to unzip {pathname} to {dir}')
        print_and_flush(str(e))
        return False

## Zip up a directory
def make_zipfile(directory, zipfile):
    absdirectory = absolute_file_name(directory)
    abszipfile = absolute_file_name(zipfile)
    # shutil.make_archive adds a .zip extension even if you already have one
    base = file_to_base_file_and_extension(abszipfile)[0]
    try:
        shutil.make_archive(base, 'zip', absdirectory)
        print_and_flush(f'--- Zipped up {directory} as {zipfile}')
        return True
    except Exception as e:
        print_and_flush(f'*** Error: something went wrong when trying to zip up {directory} as {zipfile}')
        print_and_flush(str(e))
        return False

## Save data as pickled gzipped file
def save_data_to_pickled_gzipped_file(data, pathname):
    try:
        abspathname = absolute_file_name(pathname)
        f = gzip.open(abspathname, 'wb')
        pickle.dump(data, f)
        f.close()
        print(f'--- Pickled, gzipped data written to {pathname}')
        return True
    except Exception as e:
        print_and_flush(f'*** Error: something went wrong when trying to create pickled gzipped file {pathname}')
        print_and_flush(str(e))
        return False

## Restore data from pickled gzipped file
def get_data_from_pickled_gzipped_file(pathname):
    abspathname = absolute_file_name(pathname)
    try:
        if not file_exists(abspathname):
            print_and_flush(f'*** Error: unable to find {pathname}')
            return False
        f = gzip.open(abspathname, 'rb')
        data = pickle.load(f)
        f.close()
        print(f'--- Read pickled, gzipped data from {pathname}')
        return data
    except Exception as e:
        print_and_flush(f'*** Error: something went wrong when trying to read pickled gzipped file {pathname}')
        print_and_flush(str(e))
        return False

## Remove duplicates from a list
def remove_duplicates(List):
    return list(dict.fromkeys(List))

## Make list into a multiset of the form [ ( Elt1, Freq1), ( Elt2, Freq2 ) ... ] in descending frequency order
def list_to_ordered_multiset(List):
    if not isinstance(List, list):
        utils.print_and_flush(f'*** Bad call to list_to_ordered_multiset, argument is not a list')
        return False
    Dict = {}
    for X in List:
        Key = str(X)
        if Key in Dict:
            ( Elt, Count ) = Dict[Key]
            Dict[Key] = ( Elt, Count + 1 )
        else:
            Dict[Key] = ( X, 1 )
    UnorderedMultiset = [ Dict[Key] for Key in Dict ]
    return sorted(UnorderedMultiset, key=lambda X: X[1], reverse=True)

## Slower version for lists with non-hashable elements
def remove_duplicates_general(List):
    ( Out, Dict ) = ( [], {} )
    for X in List:
        Str = str(X)
        if not Str in Dict:
            Out += [ X ]
            Dict[Str] = True
    return Out    

## Add 1 to 'Key' entry in Dict if it's there, else set it to 1
def inc_assoc(Dict, Key):
    if Key in Dict:
        Dict[Key] += 1
    else:
        Dict[Key] = 1

def inc_assoc_by_amount(Dict, Key, Amount):
    if Key in Dict:
        Dict[Key] += Amount
    else:
        Dict[Key] = Amount

## Merge two dictionaries
def merge_dicts(X, Y):
    return { **X, **Y } 

## Try to convert a string to an int, return False if not possible
def safe_string_to_int(Str):
    try:
        return int(Str, 10)
    except:
        return False

## Try to convert a string to a number, return False if not possible
def safe_string_to_number(Str):
    try:
        return int(Str, 10)
    except:
        try:
            return float(Str)
        except:
            return False

## Members of list not equal to False
def non_false_members(List):
    if not List:
        return List
    return [ X for X in List if X != False ]

## Is a list with N elements
def is_n_item_list(List, N):
    return isinstance(List, list) and len(List) == N

def is_null_string_or_spaces(Str):
    if not isinstance(Str, ( str )):
        print_and_flush(f'*** Warning: argument {Str} to is_null_string_or_spaces is not string')
        return False
    return Str == '' or Str.isspace()

## Is a list of strings
def is_list_of_strings(List):
    if not isinstance(List, list):
        return False
    for X in List:
        if not isinstance(X, str):
            return False
    return True

## Cat together a list of lists
def concatenate_lists(Lists):
    return [ X for List in Lists for X in List ]

## Prettyprint
def prettyprint(X):
    import lara_config
    if isinstance(X, lara_config.LARAConfigDict):
        X1 = X.keys_and_values()
    else:
        X1 = X
    pprint.pprint(X1)

## Get a timestamp
def timestamp():
    dateTimeObj = datetime.now()
    return dateTimeObj.strftime("%Y-%m-%d_%H-%M-%S")

## This pair of functions let you do multiple replacements efficiently
## Substitutions is a dict
def make_multiple_replacement_regex(Substitutions):
    Substrings = sorted(Substitutions, key=len, reverse=True)
    return re.compile('|'.join(map(re.escape, Substrings)))

def apply_multiple_replacement_regex(Substitutions, Regex, Str):
    return Regex.sub(lambda match: Substitutions[match.group(0)], Str)

## Split string on any one of several delimiters
def split_on_multiple_delimiters(DelimiterList, Str):
    regexPattern = '|'.join(map(re.escape, DelimiterList))
    return re.split(regexPattern, Str)

## Read a CSV file and write a JSON file consisting of the lines
def csv_to_json(CSVFile, JSONFile):
    CSVContents = read_lara_csv(CSVFile)
    if CSVContents:
        write_json_to_file(CSVContents, JSONFile)

## Read a JSON file and write a CSV file consisting of the lines if that makes sense
def json_to_csv(JSONFile, CSVFile):
    JSONContents = read_json_file(JSONFile)
    CSVContents = json_contents_to_csv_contents(JSONContents)
    if CSVContents:
        write_lara_csv(CSVContents, CSVFile)

def json_contents_to_csv_contents(JSONContents):
    if not isinstance(JSONContents, list):
        return False
    ( CSVContents, LineNumber ) = ( [], 1 )
    for Line in JSONContents:
        Line1 = json_line_to_csv_line(Line, LineNumber)
        if not Line1:
            return False
        CSVContents += [ Line1 ]
        LineNumber += 1
    return CSVContents

def json_line_to_csv_line(Line, LineNumber):
    if not isinstance(Line, list):
        print_and_flush(f'*** Error in line {LineNumber}: {Line}, not a list')
        return False
    Line1 = []
    for Element in Line:
        Element1 = json_element_to_csv_element(Element)
        if Element1 == False:
            print_and_flush(f'*** Error in line {LineNumber}: {Line} (element "{Element}")')
            return False
        Line1 += [ Element1 ]
    return Line1

def json_element_to_csv_element(Element):
    if isinstance(Element, str):
        return(Element)
    if isinstance(Element, (int, float)):
        return str(Element)
    print_and_flush(f'*** Error: {Element} is not a string or a number')
    return False

def str_to_html_str(Str):
    return Str.encode('ascii', 'xmlcharrefreplace').decode('utf-8')

def looks_like_a_url(Str):
    return isinstance(Str, str) and ( Str.startswith('https://') or Str.startswith('http://') )

# ----------------------------------

def random_choose_from_list(List):
    return random.choice(List)

def random_choose_from_weighted_list(Dict):
    if not isinstance(Dict, dict):
        print_and_flush(f'*** Error: argument to lara_utils.random_choose_from_weighted_list not a dict')
        return False
    Total = sum(Dict.values())
    Choice = random.randrange(Total)
    I = 0
    for Key in Dict:
        I += Dict[Key]
        if I >= Choice:
            return Key
    # This shouldn't happen
    print_and_flush(f'*** Error: something went wrong in lara_utils.random_choose_from_weighted_list')
    return False

def random_true(Probability):
    return random.random() > Probability
        
# ----------------------------------

def word_edit_distance(Word, Word1):
    return wagnerfischer.WagnerFischer(Word, Word1).cost

# ----------------------------------

def size_of_image(File):
    try:
        from PIL import Image
        im = Image.open(absolute_file_name(File))
        return im.size
    except Exception as e:
        print_and_flush(f'*** Error: something went wrong when trying to get size of image {Image}')
        print_and_flush(str(e))
        return False

# ----------------------------------

def length_of_mp3(File):
    try:
        from mutagen.mp3 import MP3
        audio = MP3(absolute_file_name(File))
        return audio.info.length
    except Exception as e:
        print_and_flush(f'*** Error: something went wrong when trying to get length of mp3 file {File}')
        print_and_flush(str(e))
        return False

# ----------------------------------

def is_page_tag_chunk(Chunk):
    return isinstance(Chunk, (list, tuple)) and len(Chunk) != 0 and Chunk[0] == '*page_tag*'
        
def page_tag_chunk_page_info(Chunk):
    if is_page_tag_chunk(Chunk):
        return Chunk[2]
    else:
        return False

# ----------------------------------

def get_word_pages_dir_from_params(Params):
    if Params.word_pages_directory != '':
        return absolute_file_name(Params.word_pages_directory)
    else:
        return False

def get_multimedia_dir_from_params(Params):
    WordPagesDir = get_word_pages_dir_from_params(Params)
    if WordPagesDir:
        return f'{WordPagesDir}/multimedia'
    else:
        return False

def get_corpus_dir_from_params(Params):
    if Params.corpus != '':
        return directory_for_pathname(Params.corpus)
    else:
        return False

## Add a corpus-id tag to a set of params taken from local_config
def add_corpus_id_tag_to_params(Params, Id):
    Params.corpus_id = Id
    return Params

def no_corpus_id_in_params(Params):
    Tag = Params.corpus_id
    return not Tag == 'local_files' and not isinstance(Tag, dict) and not 'corpus' in Tag

def get_corpus_id_from_params(Params):
    Tag = Params.corpus_id
    if isinstance(Tag, dict) and 'corpus' in Tag:
        return Tag['corpus']
    else:
        return Tag

def get_language_id_from_params(Params):
    Tag = Params.corpus_id
    if isinstance(Tag, dict) and 'language' in Tag:
        return Tag['language']
    else:
        return Tag

def get_l1_from_params(Params):
    Tag = Params.corpus_id
    if isinstance(Tag, dict) and 'l1' in Tag:
        return Tag['l1']
    else:
        return Tag

## For left-to-right languages the text is on the left and the word pages on the right
## For right-to-left languages, it's reversed.

def split_screen_pane_name_for_word_page_screen(Params):
    return 'aux_frame';
    # if Params.text_direction == 'rtl':
    #     return 'left'
    # else:
    #     return 'right'

def split_screen_pane_name_for_main_text_screen(Params):
    return 'main_frame';
    # if Params.text_direction == 'rtl':
    #     return 'right'
    # else:
    #     return 'left'

def relative_multimedia_dir(Params):
    return f'{Params.relative_compiled_directory}/multimedia' if Params.relative_compiled_directory != '' else 'multimedia'
